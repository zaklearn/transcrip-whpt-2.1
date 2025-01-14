import logging
from docx import Document
from io import BytesIO
from typing import Union
from datetime import datetime
def load_translations():
    """Charge les traductions pour l'interface."""
    return {
        'fr': {
            'app_title': "Transcription Audio Multilingue",
            'app_description': "Téléchargez un fichier audio pour commencer la transcription avec suivi en temps réel.",
            'settings': "Paramètres",
            'select_language': "Sélectionner la langue",
            'model_config': "Configuration du modèle",
            'select_model': "Sélectionner le modèle",
            'system_status': "État du système",
            'current_model': "Modèle actuel",
            'processing_mode': "Mode de traitement",
            'last_update': "Dernière mise à jour",
            'choose_file': "Choisir un fichier audio",
            'supported_formats': "Formats supportés : MP3, WAV, M4A",
            'start_transcription': "Démarrer la transcription",
            'processing': "Traitement en cours...",
            'transcription_completed': "Transcription terminée !",
            'success_message': "Transcription terminée avec succès ! 🎉",
            'error_message': "Une erreur s'est produite",
            'choose_format': "Choisir le format de sortie",
            'display_text': "Afficher le texte",
            'download_word': "Télécharger en Word",
            'transcription_result': "Résultat de la transcription",
            'copy_clipboard': "Copier dans le presse-papiers",
            'text_copied': "Texte copié !",
            'download_word_button': "Télécharger le document Word",
            'debug_console': "Console de débogage",
            'system_logs': "Journaux système",
            'transcription': "Transcription"
        },
        'en': {
            'app_title': "Multilingual Audio Transcription",
            'app_description': "Upload an audio file to start transcription with real-time tracking.",
            'settings': "Settings",
            'select_language': "Select Language",
            'model_config': "Model Configuration",
            'select_model': "Select Model",
            'system_status': "System Status",
            'current_model': "Current Model",
            'processing_mode': "Processing Mode",
            'last_update': "Last Update",
            'choose_file': "Choose an audio file",
            'supported_formats': "Supported formats: MP3, WAV, M4A",
            'start_transcription': "Start Transcription",
            'processing': "Processing...",
            'transcription_completed': "Transcription completed!",
            'success_message': "Transcription completed successfully! 🎉",
            'error_message': "An error occurred",
            'choose_format': "Choose output format",
            'display_text': "Display Text",
            'download_word': "Download as Word",
            'transcription_result': "Transcription Result",
            'copy_clipboard': "Copy to Clipboard",
            'text_copied': "Text copied!",
            'download_word_button': "Download Word Document",
            'debug_console': "Debug Console",
            'system_logs': "System Logs",
            'transcription': "Transcription"
        }
    }
def setup_logging() -> logging.Logger:
    """Configure logging for the application."""
    logger = logging.getLogger("WhisperApp")
    logger.setLevel(logging.INFO)
    
    handler = logging.StreamHandler()
    handler.setLevel(logging.INFO)
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    
    return logger

def create_word_document(text: str, title: str = "Transcription") -> BytesIO:
    """Create a Word document from the transcribed text."""
    doc = Document()
    
    # Add metadata
    doc.core_properties.author = "Whisper Transcription App"
    doc.core_properties.created = datetime.now()
    
    # Add content
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    
    # Add footer with timestamp
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].text = f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
